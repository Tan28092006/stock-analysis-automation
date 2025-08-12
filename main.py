import os
import pandas as pd
from datetime import datetime, timedelta
import json
import glob
from docx import Document
from groq import Groq
import yagmail
from vnstock import Vnstock

# ==== Thiết lập biến môi trường ====
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise Exception("Bạn chưa set biến môi trường GROQ_API_KEY")

client = Groq(api_key=GROQ_API_KEY)
vn = Vnstock()

# ==== Hàm tính chỉ báo ====
def EMA(series, period):
    return series.ewm(span=period, adjust=False).mean()

def SMA(series, period):
    return series.rolling(window=period).mean()

def MACD(series, fast=12, slow=26, signal=9):
    ema_fast = EMA(series, fast)
    ema_slow = EMA(series, slow)
    macd = ema_fast - ema_slow
    signal_line = EMA(macd, signal)
    hist = macd - signal_line
    return macd, signal_line, hist

def RSI(series, period=14):
    delta = series.diff()
    gain = delta.clip(lower=0)
    loss = -delta.clip(upper=0)
    avg_gain = gain.rolling(window=period).mean()
    avg_loss = loss.rolling(window=period).mean()
    rs = avg_gain / avg_loss
    rsi = 100 - (100 / (1 + rs))
    return rsi

def BollingerBands(series, period=20, num_std=2):
    mid = SMA(series, period)
    std = series.rolling(window=period).std()
    upper = mid + num_std * std
    lower = mid - num_std * std
    return upper, mid, lower

def AO(high, low, short=5, long=34):
    median_price = (high + low) / 2
    return SMA(median_price, short) - SMA(median_price, long)

# ==== Lấy top 5 cổ phiếu tăng trưởng mạnh nhất VN30 trong 14 ngày ====
def get_top_gainers_vnstock(days=14, top_n=5):
    vn30_df = vn.index_components(index_code="VN30")
    vn30_tickers = vn30_df['ticker'].tolist()

    end_date = datetime.now()
    start_date = end_date - timedelta(days=days)
    change_data = []

    for ticker in vn30_tickers:
        try:
            df = vn.stock(symbol=ticker).quote.history(
                start=start_date.strftime("%Y-%m-%d"),
                end=end_date.strftime("%Y-%m-%d"),
                interval="1d"
            )
            if df.empty:
                continue
            start_price = df['close'].iloc[0]
            end_price = df['close'].iloc[-1]
            pct_change = ((end_price - start_price) / start_price) * 100
            change_data.append((ticker, pct_change))
        except Exception as e:
            print(f"Lỗi lấy dữ liệu {ticker}: {e}")

    df_change = pd.DataFrame(change_data, columns=["Ticker", "Change"])
    df_change = df_change.sort_values(by="Change", ascending=False)
    return df_change.head(top_n)["Ticker"].tolist()

# ==== Chạy lấy danh sách top tăng ====
tickers = get_top_gainers_vnstock()
print("📈 Top 5 cổ phiếu VN30 tăng mạnh nhất 14 ngày qua:", tickers)

# ==== Thư mục lưu file ====
save_path = "outputs"
os.makedirs(save_path, exist_ok=True)

# ==== Ngày lấy dữ liệu ====
end_date = datetime.today()
start_date = end_date - timedelta(days=28)

# ==== Tải dữ liệu, tính chỉ báo, lưu JSON ====
for ticker in tickers:
    df = vn.stock(symbol=ticker).quote.history(
        start=start_date.strftime("%Y-%m-%d"),
        end=end_date.strftime("%Y-%m-%d"),
        interval="1d"
    )
    if df.empty:
        continue

    df.rename(columns={
        'time': 'Date',
        'close': 'Close',
        'open': 'Open',
        'high': 'High',
        'low': 'Low',
        'volume': 'Volume'
    }, inplace=True)

    df["EMA20"] = EMA(df["Close"], 20)
    df["EMA50"] = EMA(df["Close"], 50)
    df["EMA100"] = EMA(df["Close"], 100)
    macd, signal, hist = MACD(df["Close"])
    df["MACD"] = macd
    df["Signal"] = signal
    df["MACD_Hist"] = hist
    df["RSI"] = RSI(df["Close"])
    upper, mid, lower = BollingerBands(df["Close"])
    df["BB_Upper"] = upper
    df["BB_Mid"] = mid
    df["BB_Lower"] = lower
    df["AO"] = AO(df["High"], df["Low"])

    vol_mean = df["Volume"].rolling(window=20).mean()
    df["Breakout"] = df["Volume"] > vol_mean * 1.5

    df_reset = df.reset_index(drop=True)
    df_reset["Date"] = pd.to_datetime(df_reset["Date"]).dt.strftime("%Y-%m-%d")
    df_json_ready = df_reset.where(pd.notnull(df_reset), None)

    file_json = os.path.join(save_path, f"{ticker}.json")
    with open(file_json, "w") as f:
        json.dump(df_json_ready.to_dict(orient="records"), f, indent=2)

print("✅ Đã lưu tất cả dữ liệu JSON vào thư mục outputs!")

# ==== Hàm gọi API Groq phân tích ====
def analyze_data_with_groq(json_data):
    prompt = (
        "Bạn là chuyên gia phân tích kỹ thuật chứng khoán top 0,1%.\n"
        "Dưới đây là dữ liệu kỹ thuật của cổ phiếu (24 ngày gần nhất), "
        "hãy phân tích, nhận định xu hướng, điểm mua/bán, cảnh báo breakout, "
        "và đưa ra khuyến nghị ngắn gọn cho toàn bộ danh sách.\n\n"
        f"Dữ liệu: {json_data}\n\n"
        "Phân tích chi tiết:"
    )
    completion = client.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_completion_tokens=3000,
        top_p=1,
        reasoning_effort="medium"
    )
    return completion.choices[0].message.content

# ==== Phân tích và gom kết quả ====
json_files = glob.glob(os.path.join(save_path, "*.json"))
report_text = "Báo cáo phân tích kỹ thuật chứng khoán tự động:\n\n"

for file_path in json_files:
    with open(file_path, "r") as f:
        data_json = f.read()

    ticker_name = os.path.basename(file_path).replace(".json", "")
    print(f"Đang phân tích {ticker_name} ...")

    try:
        analysis_text = analyze_data_with_groq(data_json)
    except Exception as e:
        print(f"❌ Lỗi khi gọi API phân tích {ticker_name}: {e}")
        analysis_text = "Không có dữ liệu phân tích do lỗi API."

    report_text += f"--- Phân tích {ticker_name} ---\n{analysis_text}\n\n"
    os.remove(file_path)

# ==== Xuất file DOCX ====
doc = Document()
doc.add_heading("Báo cáo Phân tích Chỉ báo Kỹ thuật Cổ phiếu", level=1)
for line in report_text.strip().split('\n'):
    doc.add_paragraph(line)

report_path = os.path.join(save_path, "Bao_cao_phan_tich_co_phieu.docx")
doc.save(report_path)

# ==== Gửi email ====
def send_email_report(receiver_email, subject, content, attachment_path):
    sender_email = os.getenv("EMAIL_USER")
    sender_password = os.getenv("EMAIL_PASS")
    if not sender_email or not sender_password:
        raise Exception("Chưa set EMAIL_USER và EMAIL_PASS trong secrets")

    yag = yagmail.SMTP(user=sender_email, password=sender_password)
    yag.send(to=receiver_email, subject=subject, contents=content, attachments=attachment_path)
    print(f"📧 Đã gửi báo cáo tới {receiver_email}")

print(f"✅ Đã lưu báo cáo tại {report_path}")

try:
    send_email_report(
        receiver_email="vanheminhtan@gmail.com",
        subject="Báo cáo phân tích chứng khoán tự động",
        content=report_text,
        attachment_path=report_path
    )
except Exception as e:
    print(f"❌ Lỗi gửi email: {e}")
